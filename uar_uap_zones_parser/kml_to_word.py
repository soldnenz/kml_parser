import xml.etree.ElementTree as ET
import math
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
import os

def decimal_to_dms(decimal, is_latitude=True):
    """
    Преобразует десятичные координаты в формат градусы-минуты-секунды
    """
    # Определяем знак
    is_positive = decimal >= 0
    decimal = abs(decimal)
    
    # Вычисляем градусы, минуты, секунды
    degrees = int(decimal)
    minutes_float = (decimal - degrees) * 60
    minutes = int(minutes_float)
    seconds = int((minutes_float - minutes) * 60)
    
    # Форматируем строку в зависимости от типа координаты (широта/долгота)
    if is_latitude:
        direction = "N" if is_positive else "S"
        return f"{direction}{degrees:02d}{minutes:02d}{seconds:02d}"
    else:
        direction = "E" if is_positive else "W"
        return f"{direction}{degrees:03d}{minutes:02d}{seconds:02d}"

def analyze_shape(coordinates, shape_type_hint=None):
    """
    Анализирует форму на основе координат
    Возвращает тип формы и соответствующие данные
    
    shape_type_hint: подсказка типа (polygon или line)
    """
    # Если совсем мало точек, это точка или линия
    if len(coordinates) < 4:
        return "point", coordinates
    
    # Специальный случай: LineString с более чем 50 точками почти всегда круг
    # (как Ханшатыр и Флаг в примере)
    if len(coordinates) > 50 and shape_type_hint == "line":
        # Проверяем, замкнута ли линия (начало и конец совпадают)
        is_closed = (math.sqrt((coordinates[0][0] - coordinates[-1][0])**2 + 
                              (coordinates[0][1] - coordinates[-1][1])**2) < 0.0001)
        
        if is_closed:
            # Находим центр и радиус
            center_lon = sum(coord[0] for coord in coordinates) / len(coordinates)
            center_lat = sum(coord[1] for coord in coordinates) / len(coordinates)
            
            # Вычисляем среднее расстояние до центра (радиус)
            distances = []
            for lon, lat, _ in coordinates:
                distance = math.sqrt((lon - center_lon)**2 + (lat - center_lat)**2)
                distances.append(distance)
            
            avg_distance = sum(distances) / len(distances)
            radius_meters = avg_distance * 111000  # примерно метров в градусе
            
            # Для таких случаев как Ханшатыр и Флаг возвращаем круг
            return "circle", (center_lat, center_lon, radius_meters)
    
    # Проверка замкнутости фигуры (первая и последняя точки совпадают)
    is_closed = (math.sqrt((coordinates[0][0] - coordinates[-1][0])**2 + 
                          (coordinates[0][1] - coordinates[-1][1])**2) < 0.0001)
    
    # Находим центр многоугольника (центр тяжести)
    center_lon = sum(coord[0] for coord in coordinates) / len(coordinates)
    center_lat = sum(coord[1] for coord in coordinates) / len(coordinates)
    
    # Вычисляем расстояние от центра до каждой точки
    distances = []
    for lon, lat, _ in coordinates:
        distance = math.sqrt((lon - center_lon)**2 + (lat - center_lat)**2)
        distances.append(distance)
    
    # Если все расстояния примерно одинаковые, то фигура похожа на круг
    avg_distance = sum(distances) / len(distances)
    distance_variance = sum((d - avg_distance)**2 for d in distances) / len(distances)
    
    # Относительная дисперсия (отклонение / среднее)
    rel_variance = distance_variance / (avg_distance**2) if avg_distance > 0 else float('inf')
    
    # Проверка на круг для других случаев
    if (is_closed and len(coordinates) > 10) or shape_type_hint == "polygon":
        if rel_variance < 0.005:  # 0.5% для обычного Polygon
            radius_meters = avg_distance * 111000
            return "circle", (center_lat, center_lon, radius_meters)
    
    # Для простых многоугольников (мало точек)
    if len(coordinates) < 10:
        return "polygon", coordinates
    
    # Если это прямоугольник (примерно 5 точек с повтором первой) с параллельными сторонами
    if len(coordinates) == 5 and is_closed:
        sides = []
        for i in range(4):
            next_i = (i + 1) % 4
            dx = coordinates[next_i][0] - coordinates[i][0]
            dy = coordinates[next_i][1] - coordinates[i][1]
            sides.append((dx, dy))
        
        # Проверяем параллельность противоположных сторон
        parallel_1 = abs(sides[0][0] * sides[2][1] - sides[0][1] * sides[2][0]) < 0.00001
        parallel_2 = abs(sides[1][0] * sides[3][1] - sides[1][1] * sides[3][0]) < 0.00001
        
        if parallel_1 and parallel_2:
            return "rectangle", coordinates[:4]
    
    # Упрощаем сложные многоугольники и линии, сохраняя ключевые точки
    if len(coordinates) > 10:
        # Динамически подбираем шаг, чтобы получить примерно 8-12 точек
        target_points = 10
        step = max(1, len(coordinates) // target_points)
        simple_coords = coordinates[::step]
        
        # Добавляем последнюю точку, если её еще нет и фигура замкнутая
        if is_closed and simple_coords[-1] != coordinates[-1]:
            simple_coords.append(coordinates[-1])
            
        return "complex_polygon" if is_closed else "path", simple_coords
    
    return "polygon" if is_closed else "path", coordinates

def parse_kml_to_word(kml_file, output_file):
    """
    Преобразует KML файл в документ Word с координатами в формате DMS
    """
    # Разбор KML файла
    tree = ET.parse(kml_file)
    root = tree.getroot()
    
    # Создаем документ Word
    doc = Document()
    
    # Настройка полей страницы
    sections = doc.sections
    for section in sections:
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
    
    # Добавляем заголовок документа
    title = doc.add_heading('Зоны и координаты', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Добавляем краткое описание
    doc.add_paragraph('Документ содержит информацию о зонах и их географических координатах в формате DMS.')
    
    # Создаем оглавление (простой список зон)
    doc.add_heading('Список зон:', level=1)
    
    # Находим все Placemark в KML файле
    ns = {'kml': 'http://www.opengis.net/kml/2.2'}
    
    # Находим все Placemark в папке
    folder = root.find('.//kml:Folder', ns)
    if folder is None:
        folder = root  # Если нет папки, ищем во всем документе
        
    placemarks = folder.findall('.//kml:Placemark', ns)
    
    # Сначала собираем список всех зон для оглавления
    for i, placemark in enumerate(placemarks, 1):
        name_elem = placemark.find('.//kml:name', ns)
        if name_elem is None or name_elem.text is None:
            continue
        zone_name = name_elem.text
        doc.add_paragraph(f"{i}. {zone_name}", style='List Number')
    
    # Добавляем разделитель
    doc.add_page_break()
    
    # Теперь обрабатываем каждую зону подробно
    for placemark in placemarks:
        # Извлекаем имя места
        name_elem = placemark.find('.//kml:name', ns)
        if name_elem is None or name_elem.text is None:
            continue
            
        zone_name = name_elem.text
        
        # Добавляем название зоны в документ
        heading = doc.add_heading(f'Зона: {zone_name}', level=1)
        
        # Получаем координаты
        polygon = placemark.find('.//kml:Polygon/kml:outerBoundaryIs/kml:LinearRing/kml:coordinates', ns)
        line_string = placemark.find('.//kml:LineString/kml:coordinates', ns)
        
        if polygon is not None:
            coordinates_text = polygon.text
            shape_type_hint = "polygon"
        elif line_string is not None:
            coordinates_text = line_string.text
            shape_type_hint = "line"
        else:
            doc.add_paragraph('Координаты не найдены')
            continue
            
        # Преобразуем строку координат в список кортежей (lon, lat, alt)
        coordinates = []
        for coord_str in coordinates_text.strip().split():
            parts = coord_str.split(',')
            if len(parts) >= 2:
                lon = float(parts[0])
                lat = float(parts[1])
                alt = float(parts[2]) if len(parts) > 2 else 0
                coordinates.append((lon, lat, alt))
        
        # Анализируем форму на основе координат с учетом подсказки типа
        shape_type, shape_data = analyze_shape(coordinates, shape_type_hint)
        
        # Отладочная информация
        print(f"Зона: {zone_name}, Тип: {shape_type}, Точек: {len(coordinates)}")
        
        # Обрабатываем в зависимости от типа фигуры
        if shape_type == "circle":
            center_lat, center_lon, radius = shape_data
            
            # Преобразуем координаты центра в формат DMS
            lat_dms = decimal_to_dms(center_lat, True)
            lon_dms = decimal_to_dms(center_lon, False)
            
            # Добавляем информацию о круге в документ
            p = doc.add_paragraph()
            p.add_run('Тип: ').bold = True
            p.add_run('Круг')
            
            p = doc.add_paragraph()
            p.add_run('Центр: ').bold = True
            p.add_run(f'{lat_dms} {lon_dms}')
            
            p = doc.add_paragraph()
            p.add_run('Радиус: ').bold = True
            p.add_run(f'{radius:.0f} м')
        
        else:
            # Создаем таблицу для координат
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            
            # Заголовки таблицы
            header_cells = table.rows[0].cells
            header_cells[0].text = '№'
            header_cells[1].text = 'Широта (DMS)'
            header_cells[2].text = 'Долгота (DMS)'
            
            # Форматируем заголовок
            for cell in header_cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            # Заполняем таблицу координатами
            coords_to_show = shape_data
            
            for i, (lon, lat, _) in enumerate(coords_to_show, 1):
                lat_dms = decimal_to_dms(lat, True)
                lon_dms = decimal_to_dms(lon, False)
                
                row_cells = table.add_row().cells
                row_cells[0].text = str(i)
                row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row_cells[1].text = lat_dms
                row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row_cells[2].text = lon_dms
                row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Добавляем разделитель между зонами
        doc.add_paragraph('')
        
        # Если это не последняя зона, добавляем разрыв страницы
        if placemark != placemarks[-1]:
            doc.add_page_break()
    
    # Сохраняем документ
    doc.save(output_file)
    print(f'Документ сохранен: {output_file}')
    print(f'Всего зон обработано: {len(placemarks)}')

if __name__ == "__main__":
    # Путь к KML файлу и выходному файлу Word
    kml_file = "amir_zones.kml"
    output_file = "zones_coordinates.docx"
    
    # Конвертируем KML в Word
    parse_kml_to_word(kml_file, output_file) 